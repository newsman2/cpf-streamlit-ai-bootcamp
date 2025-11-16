import os
import tempfile
from typing import List

import streamlit as st
from docx import Document
from langchain.tools import ToolRuntime, tool


# Access the current conversation state
@tool
def summarize_conversation(runtime: ToolRuntime) -> str:
    """Summarize the conversation so far."""
    messages = runtime.state["messages"]

    human_msgs = sum(1 for m in messages if m.__class__.__name__ == "HumanMessage")
    ai_msgs = sum(1 for m in messages if m.__class__.__name__ == "AIMessage")
    tool_msgs = sum(1 for m in messages if m.__class__.__name__ == "ToolMessage")

    return f"Conversation has {human_msgs} user messages, {ai_msgs} AI responses, and {tool_msgs} tool results"


@tool("create_word_doc")
def create_word_doc(lines: List[str]) -> str:
    """
    Create a Word (.docx) document containing the provided lines (list of strings).
    Returns the filesystem path to the saved .docx file.
    """
    doc = Document()
    doc.add_heading("Conversation Export", level=1)
    for idx, line in enumerate(lines, start=1):
        # you can add formatting, or parse 'role:' prefix if you like
        doc.add_paragraph(f"{idx}. {line}")
    tmpdir = tempfile.gettempdir()
    filename = os.path.join(tmpdir, "conversation_export.docx")
    doc.save(filename)

    return filename
